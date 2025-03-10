from time import localtime
from anyio import current_time
from django.shortcuts import render



from django.urls import reverse_lazy
from django.contrib import messages
from django.shortcuts import render
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from .models import Nivel
from .forms import NivelForm, UserForm
####################################################################################################
from django.shortcuts import render
from django.views.generic import ListView, CreateView
from django.urls import reverse_lazy
from django.contrib import messages
from .models import Nivel
from .forms import NivelForm
from django.core.paginator import Paginator

from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib import messages

class NivelListView(LoginRequiredMixin, ListView):
    model = Nivel
    template_name = 'Nivel_gestion_alumnos/nivel_list.html'
    context_object_name = 'niveles'
    paginate_by = 10  # Paginar los niveles para no mostrar todos en una sola página
    
    def get_queryset(self):
        return Nivel.objects.all().order_by('-fecha_creacion')  # Orden por fecha de creación

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['paginator'] = Paginator(Nivel.objects.all(), self.paginate_by)
        return context

    def dispatch(self, request, *args, **kwargs):
        # Todos los grupos pueden ver los niveles
        if not request.user.groups.filter(name__in=['Administrador', 'Alumnos', 'Docente', 'Profesor']).exists():
            messages.error(request, 'No tienes permisos para ver los niveles.')
            return redirect('home')
        return super().dispatch(request, *args, **kwargs)


class NivelCreateView(LoginRequiredMixin, CreateView):
    model = Nivel
    form_class = NivelForm
    template_name = 'Nivel_gestion_alumnos/nivel_form.html'
    success_url = reverse_lazy('nivel_list')

    def form_valid(self, form):
        form.instance.creado_por = self.request.user  # Asignar el usuario autenticado
        messages.success(self.request, 'Nivel creado exitosamente.')
        return super().form_valid(form)

    def form_invalid(self, form):
        messages.error(self.request, 'Hubo un error al crear el nivel. Por favor, revisa los campos.')
        return super().form_invalid(form)

    def dispatch(self, request, *args, **kwargs):
        # Solo el Administrador puede crear niveles
        if not request.user.groups.filter(name='Administrador').exists():
            messages.error(request, 'No tienes permisos para crear niveles.')
            return redirect('nivel_list')
        return super().dispatch(request, *args, **kwargs)


class NivelUpdateView(LoginRequiredMixin, UpdateView):
    model = Nivel
    form_class = NivelForm
    template_name = 'Nivel_gestion_alumnos/nivel_form.html'
    success_url = reverse_lazy('nivel_list')

    def form_valid(self, form):
        messages.success(self.request, 'Nivel actualizado exitosamente.')
        return super().form_valid(form)
    
    def dispatch(self, request, *args, **kwargs):
        # Solo el Administrador puede actualizar niveles
        if not request.user.groups.filter(name='Administrador').exists():
            messages.error(request, 'No tienes permisos para editar niveles.')
            return redirect('nivel_list')
        return super().dispatch(request, *args, **kwargs)


class NivelDeleteView(LoginRequiredMixin, DeleteView):
    model = Nivel
    template_name = 'Nivel_gestion_alumnos/nivel_confirm_delete.html'
    success_url = reverse_lazy('nivel_list')

    def delete(self, request, *args, **kwargs):
        messages.success(self.request, 'Nivel eliminado correctamente.')
        return super().delete(request, *args, **kwargs)

    def dispatch(self, request, *args, **kwargs):
        # Solo el Administrador puede eliminar niveles
        if not request.user.groups.filter(name='Administrador').exists():
            messages.error(request, 'No tienes permisos para eliminar niveles.')
            return redirect('nivel_list')
        return super().dispatch(request, *args, **kwargs)


####################################################################################################
from django.shortcuts import render, get_object_or_404, redirect
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from .models import Grado
from .forms import GradoForm

# LISTAR GRADOS
class GradoListView(LoginRequiredMixin, ListView):
    model = Grado
    template_name = 'Grados_mi_app/grado_list.html'
    context_object_name = 'grados'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['grados'] = Grado.objects.all()
        return context
    
    def dispatch(self, request, *args, **kwargs):
        # Todos los usuarios pueden ver la lista de grados
        if not request.user.groups.filter(name__in=['Administrador', 'Alumnos', 'Docente', 'Profesor']).exists():
            messages.error(request, 'No tienes permisos para ver los grados.')
            return redirect('home')
        return super().dispatch(request, *args, **kwargs)

from django.contrib import messages
from django.urls import reverse_lazy
from django.views.generic.edit import CreateView, UpdateView, DeleteView
from .forms import GradoForm
from .models import Grado

from django.contrib import messages
from django.urls import reverse_lazy
from django.views.generic import ListView
from django.views.generic.edit import CreateView, UpdateView, DeleteView
from django.contrib.auth.mixins import LoginRequiredMixin
from django.shortcuts import redirect
from .forms import GradoForm
from .models import Grado

# LISTAR GRADOS
class GradoListView(LoginRequiredMixin, ListView):
    model = Grado
    template_name = 'Grados_mi_app/grado_list.html'
    context_object_name = 'grados'

    def get_queryset(self):
        return Grado.objects.all()

    def dispatch(self, request, *args, **kwargs):
        # Todos los usuarios en los grupos especificados pueden ver la lista de grados
        if not request.user.groups.filter(name__in=['Administrador', 'Alumnos', 'Docente', 'Profesor']).exists():
            messages.error(request, 'No tienes permisos para ver los grados.')
            return redirect('home')
        return super().dispatch(request, *args, **kwargs)

# Clase base para verificar permisos de administrador
class AdminRequiredMixin(LoginRequiredMixin):
    def tiene_permiso_admin(self, request):
        if not request.user.groups.filter(name='Administrador').exists():
            messages.error(request, 'No tienes permisos para realizar esta acción.')
            return redirect('grado_list')
        return None

# CREAR GRADO
class GradoCreateView(AdminRequiredMixin, CreateView):
    model = Grado
    form_class = GradoForm
    template_name = 'Grados_mi_app/grado_form.html'
    success_url = reverse_lazy('grado_list')

    def form_valid(self, form):
        messages.success(self.request, 'Grado creado exitosamente.')
        return super().form_valid(form)

    def form_invalid(self, form):
        messages.error(self.request, 'Hubo un error al crear el grado. Por favor, revisa los datos.')
        return super().form_invalid(form)

    def dispatch(self, request, *args, **kwargs):
        permiso = self.tiene_permiso_admin(request)
        return permiso if permiso else super().dispatch(request, *args, **kwargs)

# EDITAR GRADO
class GradoUpdateView(AdminRequiredMixin, UpdateView):
    model = Grado
    form_class = GradoForm
    template_name = 'Grados_mi_app/grado_form.html'
    success_url = reverse_lazy('grado_list')

    def dispatch(self, request, *args, **kwargs):
        permiso = self.tiene_permiso_admin(request)
        return permiso if permiso else super().dispatch(request, *args, **kwargs)

# ELIMINAR GRADO
class GradoDeleteView(AdminRequiredMixin, DeleteView):
    model = Grado
    template_name = 'Grados_mi_app/grado_confirm_delete.html'
    success_url = reverse_lazy('grado_list')

    def dispatch(self, request, *args, **kwargs):
        permiso = self.tiene_permiso_admin(request)
        return permiso if permiso else super().dispatch(request, *args, **kwargs)


####################################################################################################
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from django.contrib.auth.mixins import LoginRequiredMixin
from .models import Grado, Seccion
from .forms import GradoForm, SeccionForm

# LISTAR GRADOS
class GradoListView(LoginRequiredMixin, ListView):
    model = Grado
    template_name = 'Grados_mi_app/grado_list.html'
    context_object_name = 'grados'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['grados'] = Grado.objects.all()
        return context
    
    def dispatch(self, request, *args, **kwargs):
        if not request.user.groups.filter(name__in=['Administrador', 'Alumnos', 'Docente', 'Profesor']).exists():
            messages.error(request, 'No tienes permisos para ver los grados.')
            return redirect('home')
        return super().dispatch(request, *args, **kwargs)

class GradoCreateView(LoginRequiredMixin, CreateView):
    model = Grado
    form_class = GradoForm
    template_name = 'Grados_mi_app/grado_form.html'
    success_url = reverse_lazy('grado_list')

    def form_valid(self, form):
        messages.success(self.request, 'Grado creado exitosamente.')
        return super().form_valid(form)

    def dispatch(self, request, *args, **kwargs):
        if not request.user.groups.filter(name='Administrador').exists():
            messages.error(request, 'No tienes permisos para crear grados.')
            return redirect('grado_list')
        return super().dispatch(request, *args, **kwargs)

class GradoUpdateView(LoginRequiredMixin, UpdateView):
    model = Grado
    form_class = GradoForm
    template_name = 'Grados_mi_app/grado_form.html'
    success_url = reverse_lazy('grado_list')

    def dispatch(self, request, *args, **kwargs):
        if not request.user.groups.filter(name='Administrador').exists():
            messages.error(request, 'No tienes permisos para actualizar grados.')
            return redirect('grado_list')
        return super().dispatch(request, *args, **kwargs)

class GradoDeleteView(LoginRequiredMixin, DeleteView):
    model = Grado
    template_name = 'Grados_mi_app/grado_confirm_delete.html'
    success_url = reverse_lazy('grado_list')

    def dispatch(self, request, *args, **kwargs):
        if not request.user.groups.filter(name='Administrador').exists():
            messages.error(request, 'No tienes permisos para eliminar grados.')
            return redirect('grado_list')
        return super().dispatch(request, *args, **kwargs)

# LISTAR SECCIONES
class SeccionListView(LoginRequiredMixin, ListView):
    model = Seccion
    template_name = "seccion/seccion_list.html"
    context_object_name = "secciones"

    def dispatch(self, request, *args, **kwargs):
        if not request.user.groups.filter(name__in=['Administrador', 'Alumnos', 'Docente', 'Profesor']).exists():
            messages.error(request, 'No tienes permisos para ver las secciones.')
            return redirect('home')
        return super().dispatch(request, *args, **kwargs)

class SeccionCreateView(LoginRequiredMixin, CreateView):
    model = Seccion
    form_class = SeccionForm
    template_name = "seccion/seccion_form.html"
    success_url = reverse_lazy("seccion_list")

    def form_valid(self, form):
        messages.success(self.request, "Sección creada exitosamente.")
        return super().form_valid(form)

    def dispatch(self, request, *args, **kwargs):
        if not request.user.groups.filter(name='Administrador').exists():
            messages.error(request, 'No tienes permisos para crear secciones.')
            return redirect('seccion_list')
        return super().dispatch(request, *args, **kwargs)

class SeccionUpdateView(LoginRequiredMixin, UpdateView):
    model = Seccion
    form_class = SeccionForm
    template_name = "seccion/seccion_form.html"
    success_url = reverse_lazy("seccion_list")

    def dispatch(self, request, *args, **kwargs):
        if not request.user.groups.filter(name='Administrador').exists():
            messages.error(request, 'No tienes permisos para actualizar secciones.')
            return redirect('seccion_list')
        return super().dispatch(request, *args, **kwargs)

class SeccionDeleteView(LoginRequiredMixin, DeleteView):
    model = Seccion
    template_name = "seccion/seccion_confirm_delete.html"
    success_url = reverse_lazy("seccion_list")

    def dispatch(self, request, *args, **kwargs):
        if not request.user.groups.filter(name='Administrador').exists():
            messages.error(request, 'No tienes permisos para eliminar secciones.')
            return redirect('seccion_list')
        return super().dispatch(request, *args, **kwargs)

####################################################################################################
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from .models import Periodo
from .forms import PeriodoForm

# Vista para listar los periodos
class PeriodoListView(ListView):
    model = Periodo
    template_name = "periodo/periodo_list.html"
    context_object_name = "periodos"

# Vista para crear un nuevo periodo
class PeriodoCreateView(CreateView):
    model = Periodo
    form_class = PeriodoForm
    template_name = "periodo/periodo_form.html"
    success_url = reverse_lazy("periodo_list")

    def form_valid(self, form):
        messages.success(self.request, "Periodo creado exitosamente.")
        return super().form_valid(form)

# Vista para actualizar un periodo existente
class PeriodoUpdateView(UpdateView):
    model = Periodo
    form_class = PeriodoForm
    template_name = "periodo/periodo_form.html"
    success_url = reverse_lazy("periodo_list")

    def form_valid(self, form):
        messages.success(self.request, "Periodo actualizado exitosamente.")
        return super().form_valid(form)

# Vista para eliminar un periodo
class PeriodoDeleteView(DeleteView):
    model = Periodo
    template_name = "periodo/periodo_confirm_delete.html"
    success_url = reverse_lazy("periodo_list")

    def delete(self, request, *args, **kwargs):
        messages.success(self.request, "Periodo eliminado exitosamente.")
        return super().delete(request, *args, **kwargs)
####################################################################################################
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from .models import Alumno
from .forms import AlumnoForm

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from .models import Alumno, AlumnoDetalle
from .forms import AlumnoForm, AlumnoDetalleForm

# 📋 Listar alumnos (CBV)
from django.core.paginator import Paginator
from django.db.models import Q
from django.views.generic import ListView
from .models import Alumno

class AlumnoListView(ListView):
    model = Alumno
    template_name = "alumno/alumno_list.html"
    context_object_name = "alumnos"
    paginate_by = 10  # Número de alumnos por página

    def get_queryset(self):
        queryset = Alumno.objects.all()

        # Capturamos los parámetros de búsqueda
        query = self.request.GET.get("q", "").strip()
        seccion = self.request.GET.get("seccion", "")

        # Filtro de búsqueda por nombre, apellido o DNI
        if query:
            queryset = queryset.filter(
                Q(nombre__icontains=query) | 
                Q(apellido__icontains=query) | 
                Q(dni__icontains=query)
            )

        # Filtro por sección
        if seccion:
            queryset = queryset.filter(seccion_id=seccion)

        return queryset


# 🆕 Crear un alumno (CBV)
class AlumnoCreateView(CreateView):
    model = Alumno
    form_class = AlumnoForm
    template_name = "alumno/alumno_form.html"
    success_url = reverse_lazy("alumno_list")

    def form_valid(self, form):
        messages.success(self.request, "Alumno registrado exitosamente.")
        return super().form_valid(form)

# ✏️ Editar un alumno (CBV)
class AlumnoUpdateView(UpdateView):
    model = Alumno
    form_class = AlumnoForm
    template_name = "alumno/alumno_form.html"
    success_url = reverse_lazy("alumno_list")

    def form_valid(self, form):
        messages.success(self.request, "Datos del alumno actualizados correctamente.")
        return super().form_valid(form)

# ❌ Eliminar un alumno (CBV)
class AlumnoDeleteView(DeleteView):
    model = Alumno
    template_name = "alumno/alumno_confirm_delete.html"
    success_url = reverse_lazy("alumno_list")

    def delete(self, request, *args, **kwargs):
        messages.success(self.request, "Alumno eliminado correctamente.")
        return super().delete(request, *args, **kwargs)


# 📋 Listar detalles de alumnos (FBV)
def listar_detalles_alumnos(request):
    detalles = AlumnoDetalle.objects.select_related('alumno').all()
    return render(request, 'alumno/listar_detalles.html', {'detalles': detalles})

# 👀 Ver detalle de un alumno (FBV)
def ver_detalle_alumno(request, alumno_id):
    alumno = get_object_or_404(Alumno, id=alumno_id)
    detalle = getattr(alumno, 'detalles', None)  
    return render(request, 'alumno/detalle_alumno.html', {'alumno': alumno, 'detalle': detalle})

def crear_detalle_alumno(request, alumno_id):
    alumno = get_object_or_404(Alumno, id=alumno_id)
    
    # 🛑 Verificar si ya existe un detalle para este alumno (OneToOneField)
    if hasattr(alumno, 'detalles'):
        messages.warning(request, "El alumno ya tiene detalles registrados.")
        return redirect('ver_detalle_alumno', alumno_id=alumno.id)

    if request.method == 'POST':
        form = AlumnoDetalleForm(request.POST, request.FILES)
        if form.is_valid():
            detalle = form.save(commit=False)
            detalle.alumno = alumno  # Asignar el alumno antes de guardar
            detalle.save()
            messages.success(request, "Detalles del alumno creados con éxito.")
            return redirect('ver_detalle_alumno', alumno_id=alumno.id)
        else:
            print(form.errors)  # 📢 Mostrar errores en la consola para debug

    else:
        form = AlumnoDetalleForm()

    return render(request, 'alumno/form_detalle_alumno.html', {'form': form, 'alumno': alumno})


# ✏️ Editar detalles de un alumno (FBV)
def editar_detalle_alumno(request, alumno_id):
    alumno = get_object_or_404(Alumno, id=alumno_id)
    detalle = getattr(alumno, 'detalles', None)

    if not detalle:
        messages.warning(request, "Este alumno no tiene detalles registrados.")
        return redirect('crear_detalle_alumno', alumno_id=alumno.id)

    if request.method == 'POST':
        form = AlumnoDetalleForm(request.POST, request.FILES, instance=detalle)
        if form.is_valid():
            form.save()
            messages.success(request, "Detalles del alumno actualizados con éxito.")
            return redirect('ver_detalle_alumno', alumno_id=alumno.id)
    else:
        form = AlumnoDetalleForm(instance=detalle)

    return render(request, 'alumno/form_detalle_alumno.html', {'form': form, 'alumno': alumno})

# ❌ Eliminar detalles de un alumno (FBV)
def eliminar_detalle_alumno(request, alumno_id):
    alumno = get_object_or_404(Alumno, id=alumno_id)
    detalle = getattr(alumno, 'detalles', None)

    if not detalle:
        messages.warning(request, "Este alumno no tiene detalles registrados.")
        return redirect('listar_detalles_alumnos')

    if request.method == 'POST':
        detalle.delete()
        messages.success(request, "Detalles eliminados correctamente.")
        return redirect('listar_detalles_alumnos')

    return render(request, 'alumno/eliminar_detalle_alumno.html', {'alumno': alumno})










####################################################################################################
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from .models import Docente
from .forms import DocenteForm

from django.contrib import messages
from django.contrib.auth.mixins import UserPassesTestMixin
from django.contrib.auth.models import User
from django.shortcuts import redirect
from django.urls import reverse_lazy
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from .models import Docente
from .forms import DocenteForm


# 🔹 Mixin para verificar si el usuario es administrador
class AdminRequiredMixin(UserPassesTestMixin):
    def test_func(self):
        return self.request.user.groups.filter(name='Administrador').exists()

    def handle_no_permission(self):
        messages.error(self.request, "No tienes permisos para acceder a esta sección.")
        return redirect("docente_list")


# 🔹 Vista para listar docentes
class DocenteListView(ListView):
    model = Docente
    template_name = "docente/docente_list.html"
    context_object_name = "docentes"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["es_admin"] = self.request.user.groups.filter(name="Administrador").exists()
        return context


# 🔹 Vista para crear docentes
class DocenteCreateView(AdminRequiredMixin, CreateView):
    model = Docente
    form_class = DocenteForm
    template_name = "docente/docente_form.html"
    success_url = reverse_lazy("docente_list")

    def form_valid(self, form):
        usuario = form.cleaned_data['usuario']

        if Docente.objects.filter(usuario=usuario).exists():
            messages.error(self.request, "Este usuario ya está asignado a un docente.")
            return self.form_invalid(form)

        messages.success(self.request, "Docente registrado exitosamente.")
        return super().form_valid(form)


# 🔹 Vista para actualizar docentes
class DocenteUpdateView(AdminRequiredMixin, UpdateView):
    model = Docente
    form_class = DocenteForm
    template_name = "docente/docente_form.html"
    success_url = reverse_lazy("docente_list")

    def form_valid(self, form):
        usuario = form.cleaned_data['usuario']

        if Docente.objects.exclude(id=self.object.id).filter(usuario=usuario).exists():
            messages.error(self.request, "Este usuario ya está asignado a otro docente.")
            return self.form_invalid(form)

        messages.success(self.request, "Docente actualizado exitosamente.")
        return super().form_valid(form)


# 🔹 Vista para eliminar docentes
class DocenteDeleteView(AdminRequiredMixin, DeleteView):
    model = Docente
    template_name = "docente/docente_confirm_delete.html"
    success_url = reverse_lazy("docente_list")

    def delete(self, request, *args, **kwargs):
        messages.success(self.request, "Docente eliminado exitosamente.")
        return super().delete(request, *args, **kwargs)



####################################################################################################
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from .models import CursoAsignado
from .forms import CursoAsignadoForm

# Vista para listar los cursos asignados
class CursoAsignadoListView(ListView):
    model = CursoAsignado
    template_name = "curso_asignado/curso_asignado_list.html"
    context_object_name = "cursos_asignados"

# Vista para crear un nuevo curso asignado
class CursoAsignadoCreateView(CreateView):
    model = CursoAsignado
    form_class = CursoAsignadoForm
    template_name = "curso_asignado/curso_asignado_form.html"
    success_url = reverse_lazy("curso_asignado_list")

    def form_valid(self, form):
        messages.success(self.request, "Curso asignado exitosamente.")
        return super().form_valid(form)

# Vista para actualizar un curso asignado existente
class CursoAsignadoUpdateView(UpdateView):
    model = CursoAsignado
    form_class = CursoAsignadoForm
    template_name = "curso_asignado/curso_asignado_form.html"
    success_url = reverse_lazy("curso_asignado_list")

    def form_valid(self, form):
        messages.success(self.request, "Curso asignado actualizado exitosamente.")
        return super().form_valid(form)

# Vista para eliminar un curso asignado
class CursoAsignadoDeleteView(DeleteView):
    model = CursoAsignado
    template_name = "curso_asignado/curso_asignado_confirm_delete.html"
    success_url = reverse_lazy("curso_asignado_list")

    def delete(self, request, *args, **kwargs):
        messages.success(self.request, "Curso asignado eliminado exitosamente.")
        return super().delete(request, *args, **kwargs)

####################################################################################################
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from .models import Curso
from .forms import CursoForm

from django.urls import reverse_lazy
from django.contrib import messages
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from .models import Curso
from .forms import CursoForm

# 📌 Vista para listar los cursos
class CursoListView(ListView):
    model = Curso
    template_name = "curso/curso_list.html"
    context_object_name = "cursos"

# 📌 Vista para crear un nuevo curso
class CursoCreateView(CreateView):
    model = Curso
    form_class = CursoForm
    template_name = "curso/curso_form.html"
    success_url = reverse_lazy("curso_list")

    def form_valid(self, form):
        messages.success(self.request, "✅ Curso creado exitosamente.")
        return super().form_valid(form)

    def form_invalid(self, form):
        messages.error(self.request, "❌ Hubo un error al crear el curso. Revisa los datos.")
        return super().form_invalid(form)

# 📌 Vista para actualizar un curso existente
class CursoUpdateView(UpdateView):
    model = Curso
    form_class = CursoForm
    template_name = "curso/curso_form.html"
    success_url = reverse_lazy("curso_list")

    def form_valid(self, form):
        messages.success(self.request, "✏ Curso actualizado exitosamente.")
        return super().form_valid(form)

    def form_invalid(self, form):
        messages.error(self.request, "❌ Hubo un error al actualizar el curso. Revisa los datos.")
        return super().form_invalid(form)

# 📌 Vista para eliminar un curso
class CursoDeleteView(DeleteView):
    model = Curso
    template_name = "curso/curso_confirm_delete.html"
    success_url = reverse_lazy("curso_list")

    def delete(self, request, *args, **kwargs):
        messages.warning(request, "⚠ Curso eliminado correctamente.")
        return super().delete(request, *args, **kwargs)

    
####################################################################################################
# from django.shortcuts import render, get_object_or_404, redirect
# from .models import Alumno, Asistencia, CursoAsignado, Periodo
# from datetime import date

# def pasar_asistencia(request, curso_asignado_id):
#     curso_asignado = get_object_or_404(CursoAsignado, id=curso_asignado_id)
#     alumnos = Alumno.objects.filter(seccion=curso_asignado.seccion)
#     periodo = Periodo.objects.order_by('-fecha_inicio').first()  # Último periodo académico

#     if not periodo:
#         # Si no hay periodos registrados, mostrar un mensaje de error
#         return render(request, "asistencias/error.html", {"mensaje": "No hay períodos académicos disponibles."})

#     if request.method == "POST":
#         asistencias_actualizadas = []
#         for alumno in alumnos:
#             estado = request.POST.get(f"estado_{alumno.id}", "Presente")  # Si no hay estado, por defecto "Presente"
            
#             asistencia, created = Asistencia.objects.get_or_create(
#                 alumno=alumno, 
#                 curso_asignado=curso_asignado, 
#                 fecha=date.today(),
#                 defaults={'periodo': periodo, 'estado': estado}
#             )
            
#             # Si la asistencia ya existía y el estado cambió, actualizarlo
#             if not created and asistencia.estado != estado:
#                 asistencia.estado = estado
#                 asistencia.save()
#                 asistencias_actualizadas.append(alumno.id)

#         return redirect("pasar_asistencia_list")  # Redirige correctamente

#     return render(request, "asistencias/pasar_asistencia.html", {
#         "curso_asignado": curso_asignado,
#         "alumnos": alumnos,
#         "periodo": periodo
#     })
# from django.shortcuts import render
# from .models import CursoAsignado

# def seleccionar_curso_asistencia(request):
#     cursos_asignados = CursoAsignado.objects.all()
#     return render(request, 'asistencias/seleccionar_curso.html', {'cursos_asignados': cursos_asignados})
# from django.shortcuts import render
# from .models import Asistencia

# def listar_asistencias(request):
#     asistencias = Asistencia.objects.all()
#     return render(request, 'asistencias/listar_asistencias.html', {'asistencias': asistencias})
from django.shortcuts import render, get_object_or_404, redirect
from .models import Alumno, Asistencia, CursoAsignado, Periodo
from django.utils import timezone
from datetime import date

from django.db import IntegrityError

from datetime import date
from django.utils import timezone
from django.db import IntegrityError

from django.db import IntegrityError
from django.shortcuts import render, get_object_or_404, redirect
from django.utils import timezone
from datetime import date
from .models import CursoAsignado, Alumno, Periodo, Asistencia
from .forms import AsistenciaForm

from django.shortcuts import render, get_object_or_404, redirect
from django.utils.timezone import localtime, now
from datetime import datetime
from .models import CursoAsignado, Alumno, Periodo, Asistencia

from datetime import datetime, date
from django.shortcuts import render, get_object_or_404, redirect
from django.utils.timezone import localtime, now
from .models import Asistencia, CursoAsignado, Alumno, Periodo

# def pasar_asistencia(request, curso_asignado_id):
#     curso_asignado = get_object_or_404(CursoAsignado, id=curso_asignado_id)
#     alumnos = Alumno.objects.filter(seccion=curso_asignado.seccion)
#     periodo = Periodo.objects.order_by('-fecha_inicio').first()  # Último período académico

#     if not periodo:
#         return render(request, "asistencias/error.html", {"mensaje": "No hay períodos académicos disponibles."})

#     # Validación de la fecha ingresada en GET
#     fecha_asistencia = request.GET.get('fecha')
#     try:
#         fecha_asistencia = datetime.strptime(fecha_asistencia, "%Y-%m-%d").date() if fecha_asistencia else date.today()
#     except ValueError:
#         fecha_asistencia = date.today()

#     current_time = localtime(now()).time()  # Obtiene la hora actual del servidor

#     if request.method == "POST":
#         hora_ingresada = request.POST.get("hora", None)

#         # Validar la hora ingresada por el usuario
#         try:
#             hora_ingresada = datetime.strptime(hora_ingresada, "%H:%M").time() if hora_ingresada else current_time
#         except ValueError:
#             hora_ingresada = current_time  # En caso de error, usar la hora actual

#         for alumno in alumnos:
#             estado = request.POST.get(f"estado_{alumno.id}", "Presente")

#             Asistencia.objects.update_or_create(
#                 alumno=alumno,
#                 curso_asignado=curso_asignado,
#                 fecha=fecha_asistencia,
#                 defaults={'estado': estado, 'hora': hora_ingresada, 'periodo': periodo}
#             )

#         return redirect("listar_asistencias")  # Redirige a la lista de asistencias

#     return render(request, "asistencias/pasar_asistencia.html", {
#         "curso_asignado": curso_asignado,
#         "alumnos": alumnos,
#         "periodo": periodo,
#         "fecha_asistencia": fecha_asistencia,
#         "current_time": current_time,  # Pasar la hora actual al contexto
#     })

from django.shortcuts import render, get_object_or_404, redirect
from django.utils.timezone import localtime, now
from datetime import datetime, date
from .models import CursoAsignado, Alumno, Asistencia, Periodo

def pasar_asistencia(request, curso_asignado_id):
    curso_asignado = get_object_or_404(CursoAsignado, id=curso_asignado_id)
    alumnos = Alumno.objects.filter(seccion=curso_asignado.seccion)
    periodo = Periodo.objects.order_by('-fecha_inicio').first()  # Último período académico

    if not periodo:
        return render(request, "asistencias/error.html", {"mensaje": "No hay períodos académicos disponibles."})

    # Validación de la fecha ingresada en GET
    fecha_asistencia = request.GET.get('fecha')
    try:
        fecha_asistencia = datetime.strptime(fecha_asistencia, "%Y-%m-%d").date() if fecha_asistencia else date.today()
    except ValueError:
        fecha_asistencia = date.today()

    current_time = localtime(now()).time()  # Hora actual del servidor

    if request.method == "POST":
        hora_ingresada = request.POST.get("hora", None)

        # Validar la hora ingresada por el usuario
        if hora_ingresada:
            try:
                hora_ingresada = datetime.strptime(hora_ingresada, "%H:%M").time()
            except ValueError:
                hora_ingresada = None  # Si hay error en el formato, no guardar ninguna hora

        for alumno in alumnos:
            estado = request.POST.get(f"estado_{alumno.id}", "Presente")

            Asistencia.objects.update_or_create(
                alumno=alumno,
                curso_asignado=curso_asignado,
                fecha=fecha_asistencia,
                defaults={'estado': estado, 'hora': hora_ingresada, 'periodo': periodo}
            )

        return redirect("listar_asistencias")  # Redirige a la lista de asistencias

    return render(request, "asistencias/pasar_asistencia.html", {
        "curso_asignado": curso_asignado,
        "alumnos": alumnos,
        "periodo": periodo,
        "fecha_asistencia": fecha_asistencia,
        "current_time": current_time,  # Solo para mostrar en la plantilla
    })


from django.core.paginator import Paginator
from django.shortcuts import render
from .models import CursoAsignado

from django.core.paginator import Paginator
from django.shortcuts import render
from .models import CursoAsignado, Curso, Docente, Seccion

# def seleccionar_curso_asistencia(request):
#     query_curso = request.GET.get("curso", "")
#     query_docente = request.GET.get("docente", "")
#     query_seccion = request.GET.get("seccion", "")

#     cursos_asignados = CursoAsignado.objects.all()

#     # Filtros dinámicos
#     if query_curso:
#         cursos_asignados = cursos_asignados.filter(curso__id=query_curso)
#     if query_docente:
#         cursos_asignados = cursos_asignados.filter(docente__id=query_docente)
#     if query_seccion:
#         cursos_asignados = cursos_asignados.filter(seccion__id=query_seccion)

#     # Paginación (6 cursos por página)
#     paginator = Paginator(cursos_asignados, 6)
#     page_number = request.GET.get("page")
#     cursos_paginados = paginator.get_page(page_number)

#     # Listas para Select2
#     cursos = Curso.objects.all()
#     docentes = Docente.objects.all()
#     secciones = Seccion.objects.all()

#     return render(request, 'asistencias/seleccionar_curso.html', {
#         'cursos_asignados': cursos_paginados,
#         'cursos': cursos,
#         'docentes': docentes,
#         'secciones': secciones,
#         'query_curso': query_curso,
#         'query_docente': query_docente,
#         'query_seccion': query_seccion
#     })
from django.core.paginator import Paginator
from django.shortcuts import render
from .models import CursoAsignado, Curso, Docente, Seccion, Periodo

def seleccionar_curso_asistencia(request):
    query_nivel = request.GET.get("nivel", "")
    query_grado = request.GET.get("grado", "")
    query_curso_asignado = request.GET.get("curso_asignado", "")
    query_docente = request.GET.get("docente", "")
    query_seccion = request.GET.get("seccion", "")
    query_bimestre = request.GET.get("bimestre", "")

    cursos_asignados = CursoAsignado.objects.all()

    # Filtros dinámicos
    if query_nivel:
        cursos_asignados = cursos_asignados.filter(seccion__grado__nivel__id=query_nivel)
    if query_grado:
        cursos_asignados = cursos_asignados.filter(seccion__grado__id=query_grado)
    if query_curso_asignado:
        cursos_asignados = cursos_asignados.filter(id=query_curso_asignado)
    if query_docente:
        cursos_asignados = cursos_asignados.filter(docente__id=query_docente)
    if query_seccion:
        cursos_asignados = cursos_asignados.filter(seccion__id=query_seccion)
    if query_bimestre:
        cursos_asignados = cursos_asignados.filter(asistencias__periodo__id=query_bimestre).distinct()

    # Paginación (6 cursos por página)
    paginator = Paginator(cursos_asignados, 6)
    page_number = request.GET.get("page")
    cursos_paginados = paginator.get_page(page_number)

    # Listas para Select2
    niveles = Nivel.objects.all()
    grados = Grado.objects.all()
    cursos_asignados_list = CursoAsignado.objects.all()
    docentes = Docente.objects.all()
    secciones = Seccion.objects.all()
    periodos = Periodo.objects.all()

    return render(request, 'asistencias/seleccionar_curso.html', {
        'cursos_asignados': cursos_paginados,
        'niveles': niveles,
        'grados': grados,
        'cursos_asignados_list': cursos_asignados_list,
        'docentes': docentes,
        'secciones': secciones,
        'periodos': periodos,
        'query_nivel': query_nivel,
        'query_grado': query_grado,
        'query_curso_asignado': query_curso_asignado,
        'query_docente': query_docente,
        'query_seccion': query_seccion,
        'query_bimestre': query_bimestre
    })




from django.core.paginator import Paginator
from django.utils.timezone import localdate

from django.shortcuts import render
from django.core.paginator import Paginator
from django.utils.timezone import localdate
from .models import Asistencia, Curso, Seccion
from datetime import time

from django.shortcuts import render
from django.core.paginator import Paginator
from django.utils.timezone import localdate
from .models import Asistencia, Curso, Seccion

from django.shortcuts import render
from django.core.paginator import Paginator
from django.utils.timezone import localdate
from django.contrib import messages
from .models import Asistencia, Curso, Seccion

from django.shortcuts import render
from django.utils.timezone import localdate
from django.core.paginator import Paginator
from django.contrib import messages
from django.http import HttpResponse
import openpyxl
from datetime import datetime
from .models import Asistencia, Curso, Seccion

from django.shortcuts import render
from django.utils.timezone import localdate
from datetime import datetime
from django.core.paginator import Paginator
from django.contrib import messages
from .models import Asistencia, CursoAsignado, Seccion, Nivel, Grado

from datetime import datetime
from django.utils.timezone import localdate
from django.core.paginator import Paginator
from django.contrib import messages
from django.shortcuts import render
from .models import Asistencia, CursoAsignado, Seccion, Nivel, Grado  # Importa los modelos que usas

from django.shortcuts import render
from django.utils.timezone import localdate
from django.core.paginator import Paginator
from django.contrib import messages
from datetime import datetime
from .models import Asistencia, CursoAsignado, Seccion, Nivel, Grado

def listar_asistencias(request):
    # Obtener filtros desde la URL
    fecha_filtrada = request.GET.get('fecha')
    curso_asignado_id = request.GET.get('curso_asignado')
    seccion_id = request.GET.get('seccion')
    nivel_id = request.GET.get('nivel')
    grado_id = request.GET.get('grado')

    errores = []

    # Validación de la fecha
    if fecha_filtrada:
        try:
            fecha_filtrada = datetime.strptime(fecha_filtrada, "%Y-%m-%d").date()
        except ValueError:
            errores.append("La fecha ingresada no es válida. Introduzca un formato correcto (YYYY-MM-DD).")
            fecha_filtrada = None
    else:
        fecha_filtrada = localdate()  # Usa la fecha actual por defecto

    # Filtrado de asistencias con prefetching para optimizar consultas
    asistencias = Asistencia.objects.select_related(
        'alumno', 'curso_asignado', 'curso_asignado__curso', 'curso_asignado__seccion',
        'curso_asignado__curso__grado', 'curso_asignado__curso__grado__nivel'
    )

    if fecha_filtrada:
        asistencias = asistencias.filter(fecha=fecha_filtrada)

    # Filtrar por CursoAsignado
    if curso_asignado_id and curso_asignado_id.isdigit():
        if CursoAsignado.objects.filter(id=curso_asignado_id).exists():
            asistencias = asistencias.filter(curso_asignado_id=curso_asignado_id)
        else:
            errores.append("El curso asignado seleccionado no es válido.")

    # Filtrar por Sección
    if seccion_id and seccion_id.isdigit():
        if Seccion.objects.filter(id=seccion_id).exists():
            asistencias = asistencias.filter(curso_asignado__seccion_id=seccion_id)
        else:
            errores.append("La sección seleccionada no es válida.")

    # Filtrar por Nivel
    if nivel_id and nivel_id.isdigit():
        if Nivel.objects.filter(id=nivel_id).exists():
            asistencias = asistencias.filter(curso_asignado__curso__grado__nivel_id=nivel_id)
        else:
            errores.append("El nivel seleccionado no es válido.")

    # Filtrar por Grado
    if grado_id and grado_id.isdigit():
        if Grado.objects.filter(id=grado_id).exists():
            asistencias = asistencias.filter(curso_asignado__curso__grado_id=grado_id)
        else:
            errores.append("El grado seleccionado no es válido.")

    # Ordenar por fecha (más recientes primero)
    asistencias = asistencias.order_by('-fecha')

    # Paginación: 10 asistencias por página
    paginator = Paginator(asistencias, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Obtener listas de opciones para filtros
    cursos_asignados = CursoAsignado.objects.all()
    secciones = Seccion.objects.all()
    niveles = Nivel.objects.all()
    grados = Grado.objects.all()

    # Mostrar mensajes de advertencia si hay errores
    for error in errores:
        messages.warning(request, error)

    return render(request, 'asistencias/listar_asistencias.html', {
        'asistencias': page_obj,
        'fecha_filtrada': fecha_filtrada,
        'cursos_asignados': cursos_asignados,
        'secciones': secciones,
        'niveles': niveles,
        'grados': grados,
        'curso_asignado_id': curso_asignado_id,
        'seccion_id': seccion_id,
        'nivel_id': nivel_id,
        'grado_id': grado_id,
    })

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from django.utils.timezone import localdate
from .models import Asistencia, Curso, Seccion
from .forms import AsistenciaForm

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import Asistencia
from .forms import AsistenciaForm

def actualizar_asistencia(request, asistencia_id):
    asistencia = get_object_or_404(Asistencia, id=asistencia_id)

    if request.method == "POST":
        form = AsistenciaForm(request.POST, instance=asistencia)
        if form.is_valid():
            form.save()
            messages.success(request, "✅ Asistencia actualizada correctamente.")
            return redirect('listar_asistencias')
        else:
            messages.error(request, "❌ Error al actualizar la asistencia.")
    else:
        form = AsistenciaForm(instance=asistencia)

    return render(request, 'asistencias/actualizar_asistencia.html', {'form': form, 'asistencia': asistencia})

def eliminar_asistencia(request, asistencia_id):
    asistencia = get_object_or_404(Asistencia, id=asistencia_id)

    if request.method == "POST":
        asistencia.delete()
        messages.success(request, "✅ Asistencia eliminada correctamente.")
        return redirect('listar_asistencias')

    return render(request, 'asistencias/eliminar_asistencia.html', {'asistencia': asistencia})






def exportar_asistencias(request):
    # Obtener los filtros desde la URL
    fecha = request.GET.get('fecha')
    curso_id = request.GET.get('curso')
    seccion_id = request.GET.get('seccion')

    # Validación y conversión de la fecha
    if fecha:
        try:
            fecha = datetime.strptime(fecha, "%Y-%m-%d").date()
        except ValueError:
            fecha = None  # Si la fecha no es válida, no se filtra por fecha

    # Validación de curso y sección
    curso_id = int(curso_id) if curso_id and curso_id.isdigit() else None
    seccion_id = int(seccion_id) if seccion_id and seccion_id.isdigit() else None

    # Filtrar asistencias según los parámetros proporcionados
    asistencias = Asistencia.objects.all()
    if fecha:
        asistencias = asistencias.filter(fecha=fecha)
    if curso_id:
        asistencias = asistencias.filter(curso_asignado__curso_id=curso_id)
    if seccion_id:
        asistencias = asistencias.filter(alumno__seccion_id=seccion_id)

    # Crear un archivo Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Asistencias"

    # Escribir los encabezados
    headers = ["Alumno", "Curso", "Fecha", "Estado"]
    ws.append(headers)

    # Agregar los datos de asistencia
    for asistencia in asistencias:
        ws.append([
            f"{asistencia.alumno.nombre} {asistencia.alumno.apellido}",
            asistencia.curso_asignado.curso.nombre,
            asistencia.fecha.strftime("%d/%m/%Y"),
            asistencia.estado
        ])

    # Preparar la respuesta HTTP con el archivo
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response["Content-Disposition"] = 'attachment; filename="asistencias.xlsx"'
    wb.save(response)

    return response

####################################################################################################
from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse_lazy
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from .models import Trabajo
from .forms import TrabajoForm

from django.views.generic import ListView
from django.utils.timezone import now
from django.db.models import Q
from .models import Trabajo

from django.shortcuts import render
from django.views.generic import ListView


from django.shortcuts import render
from django.views.generic import ListView
from .models import Trabajo, CursoAsignado, Docente, Periodo

from django.core.paginator import Paginator
from django.shortcuts import render
from django.views.generic import ListView
from .models import Trabajo, Docente, Curso, Periodo

from django.views.generic import ListView
from django.db.models import Q
from .models import Trabajo, Docente, Curso, Periodo

class TrabajoListView(ListView): 
    model = Trabajo
    template_name = "trabajos/trabajo_list.html"
    context_object_name = "trabajos"
    paginate_by = 10  # Paginación con 10 elementos por página

    def get_queryset(self):
        queryset = Trabajo.objects.all()
        
        # Filtrar por docente
        docente_id = self.request.GET.get("docente")
        if docente_id:
            queryset = queryset.filter(curso_asignado__docente_id=docente_id)
        
        # Filtrar por curso
        curso_id = self.request.GET.get("curso")
        if curso_id:
            queryset = queryset.filter(curso_asignado__curso_id=curso_id)
        
        # Filtrar por bimestre
        bimestre_id = self.request.GET.get("bimestre")
        if bimestre_id:
            queryset = queryset.filter(periodo_id=bimestre_id)
        
        # Filtrar por fecha de entrega
        fecha_entrega = self.request.GET.get("fecha_entrega")
        if fecha_entrega:
            queryset = queryset.filter(fecha_entrega=fecha_entrega)
        
        # Filtrar por nombre de trabajo (búsqueda)
        nombre_trabajo = self.request.GET.get("nombre_trabajo")
        if nombre_trabajo:
            queryset = queryset.filter(Q(titulo__icontains=nombre_trabajo))
        
        # Ordenar por fecha de entrega (los más próximos primero)
        queryset = queryset.order_by("fecha_entrega", "hora_entrega")
        
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["docentes"] = Docente.objects.all()
        context["cursos"] = Curso.objects.all()
        context["periodos"] = Periodo.objects.all()
        context["nombre_trabajo"] = self.request.GET.get("nombre_trabajo", "")
        return context
import openpyxl
from django.http import HttpResponse
from django.views import View
from .models import Trabajo

class ExportarTrabajosExcel(View):
    def get(self, request, *args, **kwargs):
        trabajos = Trabajo.objects.all()

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Trabajos"

        # Encabezados
        headers = ["Título", "Docente", "Curso", "Fecha de Entrega", "Hora de Entrega", "Periodo", "Tipo", "Obligatorio"]
        ws.append(headers)

        # Datos
        for trabajo in trabajos:
            ws.append([
                trabajo.titulo,
                trabajo.curso_asignado.docente.nombre if trabajo.curso_asignado.docente else "N/A",
                trabajo.curso_asignado.curso.nombre if trabajo.curso_asignado.curso else "N/A",
                trabajo.fecha_entrega.strftime("%Y-%m-%d") if trabajo.fecha_entrega else "",
                trabajo.hora_entrega.strftime("%H:%M") if trabajo.hora_entrega else "",
                trabajo.periodo.nombre if trabajo.periodo else "N/A",
                trabajo.tipo_trabajo,
                "Sí" if trabajo.es_obligatorio else "No"
            ])

        response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        response["Content-Disposition"] = 'attachment; filename="trabajos.xlsx"'
        wb.save(response)

        return response




# Crear trabajo
class TrabajoCreateView(CreateView):
    model = Trabajo
    form_class = TrabajoForm
    template_name = "trabajos/trabajo_form.html"
    success_url = reverse_lazy("trabajo_list")

# Editar trabajo
class TrabajoUpdateView(UpdateView):
    model = Trabajo
    form_class = TrabajoForm
    template_name = "trabajos/trabajo_form.html"
    success_url = reverse_lazy("trabajo_list")

# Eliminar trabajo
class TrabajoDeleteView(DeleteView):
    model = Trabajo
    template_name = "trabajos/trabajo_confirm_delete.html"
    success_url = reverse_lazy("trabajo_list")

####################################################################################################
from django.shortcuts import get_object_or_404, redirect
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from django.contrib import messages
from .models import Calificacion
from .forms import CalificacionForm

from django.db.models import Q

from django.db.models import Q
from django.views.generic import ListView
from .models import Calificacion

from django.views.generic import ListView
from django.shortcuts import render
from django.db.models import Q
from .models import Calificacion
from django.db.models import Q
from django.views.generic import ListView
from .models import Calificacion

from django.db.models import Q
from django.views.generic import ListView
from .models import Calificacion

from django.shortcuts import render
from django.views.generic import ListView
from django.db.models import Q
from .models import Calificacion, Curso, Docente

from django.db.models import Q

from django.http import HttpResponse
from django.shortcuts import render
from django.db.models import Q
import pandas as pd
from .models import Calificacion
from .models import Curso, Docente

class CalificacionListView(LoginRequiredMixin,ListView):
    model = Calificacion
    template_name = 'calificacion/calificacion_list.html'
    context_object_name = 'calificaciones'
    paginate_by = 10  # Paginación: 10 calificaciones por página

    def get_queryset(self):
        # Optimizar las consultas usando select_related para evitar consultas adicionales
        queryset = super().get_queryset().select_related(
            'trabajo', 
            'alumno', 
            'curso_asignado__curso', 
            'curso_asignado__docente'
        )

        # Obtener parámetros de búsqueda desde la solicitud GET
        trabajo = self.request.GET.get('trabajo', '')
        alumno = self.request.GET.get('alumno', '')
        estado = self.request.GET.get('estado', '')
        curso_asignado = self.request.GET.get('curso_asignado', '')
        docente = self.request.GET.get('docente', '')

        # Filtrar por trabajo
        if trabajo:
            queryset = queryset.filter(trabajo__titulo__icontains=trabajo)
        
        # Filtrar por alumno (busca en nombre o apellido)
        if alumno:
            queryset = queryset.filter(
                Q(alumno__nombre__icontains=alumno) | Q(alumno__apellido__icontains=alumno)
            )
        
        # Filtrar por estado (verifica si hay un valor y lo aplica)
        if estado:
            queryset = queryset.filter(estado=estado)

        # Filtrar por curso asignado (por nombre del curso asociado al trabajo)
        if curso_asignado:
            queryset = queryset.filter(curso_asignado__curso__nombre__icontains=curso_asignado)
        
        # Filtrar por docente (por nombre del docente asociado al curso asignado)
        if docente:
            queryset = queryset.filter(curso_asignado__docente__nombre__icontains=docente)

        return queryset


    #Roles permitidos
    def dispatch(self, request, *args, **kwargs):
        # Permitir que todos los usuarios vean la vista, pero solo ciertos grupos accedan a ella
        if not request.user.groups.filter(name__in=['Administrador', 'Profesor']).exists():
            messages.error(request, 'No tienes permisos para ver la lista de calificaciones.')
            return redirect('home')  # Redirigir a la página de inicio u otra vista permitida
            
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['trabajo_filtro'] = self.request.GET.get('trabajo', '')
        context['alumno_filtro'] = self.request.GET.get('alumno', '')
        context['estado_filtro'] = self.request.GET.get('estado', '')
        context['curso_asignado_filtro'] = self.request.GET.get('curso_asignado', '')
        context['docente_filtro'] = self.request.GET.get('docente', '')
        context['cursos'] = Curso.objects.all()  # Pasar los cursos disponibles
        context['docentes'] = Docente.objects.all()  # Pasar los docentes disponibles
        return context


def exportar_calificaciones_excel(request):
    # Obtener los filtros de búsqueda
    trabajo_filtro = request.GET.get('trabajo', '')
    alumno_filtro = request.GET.get('alumno', '')
    estado_filtro = request.GET.get('estado', '')
    curso_asignado_filtro = request.GET.get('curso_asignado', '')
    docente_filtro = request.GET.get('docente', '')

    # Filtrar los resultados según los filtros aplicados
    calificaciones = Calificacion.objects.all()

    if trabajo_filtro:
        calificaciones = calificaciones.filter(trabajo__titulo__icontains=trabajo_filtro)
    if alumno_filtro:
        calificaciones = calificaciones.filter(alumno__nombre__icontains=alumno_filtro)
    if estado_filtro:
        calificaciones = calificaciones.filter(estado=estado_filtro)
    if curso_asignado_filtro:
        calificaciones = calificaciones.filter(curso_asignado__curso__nombre__icontains=curso_asignado_filtro)
    if docente_filtro:
        calificaciones = calificaciones.filter(curso_asignado__docente__nombre__icontains=docente_filtro)

    # Crear un DataFrame con los datos filtrados
    data = []
    for calificacion in calificaciones:
        data.append({
            'Trabajo': calificacion.trabajo.titulo,
            'Alumno': f"{calificacion.alumno.nombre} {calificacion.alumno.apellido}",
            'Curso': calificacion.curso_asignado.curso.nombre,
            'Nota': calificacion.nota,
            'Estado': calificacion.estado,
            'Fecha de entrega': calificacion.trabajo.fecha_entrega,
        })

    # Crear el DataFrame
    df = pd.DataFrame(data)

    # Generar el archivo Excel
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=calificaciones.xlsx'
    
    # Convertir el DataFrame a Excel y escribirlo en la respuesta
    with pd.ExcelWriter(response, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Calificaciones')

    return response
####################################################################################################
#Configuracion Administradores
# Mixin para restringir acceso solo a administradores
from django.contrib.auth.mixins import UserPassesTestMixin, LoginRequiredMixin
class SoloAdministradoresMixin(UserPassesTestMixin):
    def test_func(self):
        return self.request.user.groups.filter(name="Administrador").exists()

    def handle_no_permission(self):
        messages.error(self.request, "🚫 No tienes permisos para realizar esta acción.")
        return redirect("calificacion_list")  # Redirige a la lista de calificaciones


####################################################################################################    
# ✅ Crear Calificacións
class CalificacionCreateView(LoginRequiredMixin, SoloAdministradoresMixin,CreateView):
    model = Calificacion
    form_class = CalificacionForm
    template_name = "calificacion/calificacion_form.html"
    success_url = reverse_lazy("calificacion_list")

    def form_valid(self, form):
        messages.success(self.request, "✅ Calificación registrada correctamente.")
        return super().form_valid(form)

# ✅ Editar Calificación
class CalificacionUpdateView(LoginRequiredMixin, SoloAdministradoresMixin,UpdateView):
    model = Calificacion
    form_class = CalificacionForm
    template_name = "calificacion/calificacion_form.html"

    def form_valid(self, form):
        messages.success(self.request, "✅ Calificación actualizada correctamente.")
        return super().form_valid(form)

    def get_success_url(self):
        return self.request.path  # 📌 Redirige a la misma página después de editar

# ✅ Eliminar Calificación (con mensaje de confirmación)
class CalificacionDeleteView(LoginRequiredMixin, SoloAdministradoresMixin,DeleteView):
    model = Calificacion
    template_name = "calificacion/calificacion_confirm_delete.html"
    success_url = reverse_lazy("calificacion_list")

    def delete(self, request, *args, **kwargs):
        messages.success(request, "❌ Calificación eliminada correctamente.")
        return super().delete(request, *args, **kwargs)



from django.shortcuts import render, get_object_or_404, redirect
from django.http import JsonResponse
from .models import Trabajo, Alumno, Calificacion
from .forms import CalificacionForm

from django.shortcuts import get_object_or_404, render, redirect
from django.contrib import messages
from decimal import Decimal
from .models import Trabajo, Alumno, Calificacion

from django.shortcuts import get_object_or_404, render, redirect
from django.contrib import messages
from decimal import Decimal
from .models import Trabajo, Alumno, Calificacion

from decimal import Decimal
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import Trabajo, Alumno, Calificacion

from decimal import Decimal, InvalidOperation
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import Trabajo, Alumno, Calificacion

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from decimal import Decimal, InvalidOperation
from .models import Trabajo, Alumno, Calificacion
from decimal import Decimal, InvalidOperation
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import Trabajo, Alumno, Calificacion

from decimal import Decimal, InvalidOperation
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import Trabajo, Alumno, Calificacion
from decimal import Decimal, InvalidOperation
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib import messages
from .models import Trabajo, Calificacion, Alumno

from django.shortcuts import render
from django.db.models import Prefetch
from .models import Trabajo, CursoAsignado, Docente, Curso, Seccion

from django.shortcuts import render
from django.core.paginator import Paginator
from django.db.models import Q
from django.utils import timezone
from .models import Trabajo, Docente, Curso, Seccion

from django.shortcuts import render
from django.core.paginator import Paginator
from django.db.models import Q
from django.utils import timezone
from django.http import JsonResponse
from .models import Trabajo, Docente, Curso, Seccion

from django.shortcuts import render
from django.db.models import Q
from django.core.paginator import Paginator
from django.utils import timezone
from .models import Trabajo, Docente, Curso, Seccion

from django.shortcuts import render
from django.utils import timezone
from django.core.paginator import Paginator
from django.db.models import Q
from .models import Trabajo, Docente, Curso, Seccion
from django.db.models import Q
from django.core.paginator import Paginator
from django.utils import timezone
from django.shortcuts import render
from gestion_alumnos.models import Trabajo, Docente, Curso, Seccion

from django.shortcuts import render
from django.db.models import Q
from django.core.paginator import Paginator
from django.utils import timezone
from .models import Trabajo, Docente, Curso, Seccion

from django.shortcuts import render
import os
from django.conf import settings

from django.shortcuts import render
import os
from django.conf import settings

import os
import random
from django.conf import settings
from django.shortcuts import render

####################################################################################################

import os
import random
from django.conf import settings
from django.shortcuts import render

import os
import random
from django.conf import settings
from django.shortcuts import render

import os
import random
from django.conf import settings
from django.shortcuts import render
import os
import random
from django.conf import settings
from django.core.paginator import Paginator
from django.shortcuts import render

from django.core.paginator import Paginator
import os
import random
from django.conf import settings
from django.shortcuts import render

import os
import random
from django.conf import settings
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.shortcuts import render

def reproductor_mp3(request):
    # Ruta completa a la carpeta de música
    media_path = os.path.join(settings.MEDIA_ROOT, "musica")

    # Crear la carpeta si no existe
    if not os.path.exists(media_path):
        os.makedirs(media_path)

    # Verificar archivos en la carpeta
    archivos = os.listdir(media_path)

    # Filtrar canciones MP3 (asegurándonos de que sea insensible a mayúsculas)
    canciones = [f"/media/musica/{archivo}" for archivo in archivos if archivo.lower().endswith(".mp3")]

    if not canciones:
        print("No se encontraron archivos MP3")

    # Obtener el filtro de búsqueda (si lo hay)
    filtro = request.GET.get('filtro', '').lower()
    if filtro:
        canciones = [cancion for cancion in canciones if filtro in cancion.lower()]

    # Mezclar canciones de manera aleatoria (opcional)
    random.shuffle(canciones)

    # Paginación de canciones (10 canciones por página)
    paginator = Paginator(canciones, 10)  # Paginación: 10 canciones por página
    page_number = request.GET.get('page', 1)  # Establecer un valor predeterminado para page

    try:
        page_obj = paginator.get_page(page_number)  # Obtiene la página solicitada
    except PageNotAnInteger:
        page_obj = paginator.get_page(1)  # Si no es un número válido, redirige a la primera página
    except EmptyPage:
        page_obj = paginator.get_page(paginator.num_pages)  # Si la página está vacía, redirige a la última página

    # Enviar las canciones paginadas a la plantilla
    return render(request, "musica/reproductor_mp3.html", {
        "page_obj": page_obj,
        "filtro": filtro
    })


def vista_niveles(request):
    return render(request, 'partials/niveles.html')

def vista_alumnos(request):
    return render(request, 'partials/alumnos.html')

####################################################################################################


# Función que verifica si el usuario es Administrador o Profesor
from django.contrib.auth.decorators import login_required, user_passes_test
def es_admin_o_profesor(user):
    return user.groups.filter(name__in=["Administrador", "Profesor"]).exists()



####################################################################################################


from django.db.models import Q
from django.shortcuts import render, redirect
from .models import Trabajo, Docente, Curso, Seccion
@login_required
@user_passes_test(es_admin_o_profesor)
def seleccionar_trabajo(request):
    print("Parámetros recibidos:", request.GET.dict())  # Ver qué llega

    trabajos = Trabajo.objects.select_related(
        'curso_asignado__docente',
        'curso_asignado__curso',
        'curso_asignado__seccion'
    ).order_by('fecha_entrega')

    # Obtener valores de filtros
    docente_id = request.GET.get('docente', '').strip()
    curso_id = request.GET.get('curso', '').strip()
    seccion_id = request.GET.get('seccion', '').strip()

    # Aplicar filtros solo si hay valores válidos
    filtros = Q()
    if docente_id and docente_id.isdigit():
        filtros &= Q(curso_asignado__docente_id=int(docente_id))
    if curso_id and curso_id.isdigit():
        filtros &= Q(curso_asignado__curso_id=int(curso_id))
    if seccion_id and seccion_id.isdigit():
        filtros &= Q(curso_asignado__seccion_id=int(seccion_id))

    if filtros:
        trabajos = trabajos.filter(filtros)

    contexto = {
        'trabajos': trabajos,
        'docentes': Docente.objects.all(),
        'cursos': Curso.objects.all(),
        'secciones': Seccion.objects.all(),
        'docente_id': docente_id,  # Para mantener el filtro al volver
        'curso_id': curso_id,      # Para mantener el filtro al volver
        'seccion_id': seccion_id   # Para mantener el filtro al volver
    }

    return render(request, 'calificacion/seleccionar_trabajo.html', contexto)

####################################################################################################
# Función que verifica si el usuario es Administrador o Profesor
from django.contrib.auth.decorators import login_required, user_passes_test
def es_admin_o_profesor(user):
    return user.groups.filter(name__in=["Administrador", "Profesor"]).exists()



####################################################################################################

@login_required
@user_passes_test(es_admin_o_profesor)
def calificar_trabajo(request, trabajo_id):
    trabajo = get_object_or_404(Trabajo, id=trabajo_id)
    
    # Obtener los alumnos de la misma sección del curso asignado al trabajo
    alumnos = Alumno.objects.filter(seccion=trabajo.curso_asignado.seccion)

    if request.method == "POST":
        errores = 0
        alumnos_seleccionados = request.POST.getlist("alumno_id")  # ✅ Esto funciona mejor


        if not alumnos_seleccionados:
            messages.warning(request, "⚠️ No seleccionaste ningún alumno.")
            return redirect("calificar_trabajo", trabajo_id=trabajo.id)

        for alumno_id in alumnos_seleccionados:
            alumno = get_object_or_404(Alumno, id=alumno_id)
            nota_str = request.POST.get(f"nota_{alumno_id}", "").replace(',', '.').strip()
            estado = request.POST.get(f"estado_{alumno_id}", "Pendiente").capitalize()
            comentarios = request.POST.get(f"comentarios_{alumno_id}", "").strip() or None

            # Validación de la nota con manejo de excepciones
            try:
                nota = Decimal(nota_str)
                if not (0 <= nota <= 20):
                    raise ValueError("Nota fuera de rango")
            except (InvalidOperation, ValueError):
                messages.error(request, f"❌ La nota de {alumno.nombre} no es válida.")
                errores += 1
                continue

            # Asegurar que el estado sea válido
            if estado not in ['Pendiente', 'Revisado', 'Aprobado']:
                estado = 'Pendiente'

            # Actualizar o crear la calificación evitando duplicados
            Calificacion.objects.update_or_create(
                trabajo=trabajo, alumno=alumno, curso_asignado=trabajo.curso_asignado,
                defaults={"nota": nota, "estado": estado, "comentarios": comentarios}
            )

        # Mensajes de éxito o advertencia
        if errores == 0:
            messages.success(request, "✅ Calificaciones guardadas correctamente.")
        else:
            messages.warning(request, f"⚠️ Hubo {errores} errores en la calificación.")

        return redirect("seleccionar_trabajo")  # Asegúrate de tener esta URL definida

    return render(request, "calificacion/calificar_trabajo.html", {"trabajo": trabajo, "alumnos": alumnos})


####################################################################################################
from django.shortcuts import render
from .models import EventoCalendario

def calendario_academico(request):
    eventos = EventoCalendario.objects.all()
    return render(request, 'calendario/calendario.html', {'eventos': eventos})
####################################################################################################
from django.shortcuts import render, redirect, get_object_or_404
from .models import EventoCalendario
from .forms import EventoCalendarioForm
from django.contrib import messages

from django.core.paginator import Paginator
from django.db.models import Q
from django.shortcuts import render
from .models import EventoCalendario

def calendario_academico(request):
    # Obtener los parámetros de filtro desde la URL
    tipo_evento = request.GET.get('tipo_evento', '')
    fecha_inicio = request.GET.get('fecha_inicio', '')
    fecha_fin = request.GET.get('fecha_fin', '')

    # Filtrar eventos según los parámetros
    eventos = EventoCalendario.objects.all()
    
    if tipo_evento:
        eventos = eventos.filter(tipo_evento=tipo_evento)

    if fecha_inicio and fecha_fin:
        eventos = eventos.filter(fecha_inicio__gte=fecha_inicio, fecha_fin__lte=fecha_fin)

    # Paginación (5 eventos por página)
    paginator = Paginator(eventos, 5)  
    page_number = request.GET.get('page')
    eventos_paginados = paginator.get_page(page_number)

    return render(request, 'calendario/calendario.html', {
        'eventos': eventos_paginados,
        'tipo_evento': tipo_evento,
        'fecha_inicio': fecha_inicio,
        'fecha_fin': fecha_fin
    })

from django.http import JsonResponse

from django.shortcuts import render
from django.contrib import messages
from .forms import EventoCalendarioForm

def agregar_evento(request):
    form = EventoCalendarioForm()

    if request.method == 'POST':
        form = EventoCalendarioForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "✅ Evento agregado exitosamente.")
            form = EventoCalendarioForm()  # Reiniciar el formulario después de guardar

    return render(request, 'calendario/evento_form.html', {'form': form, 'accion': 'Agregar Evento'})


from django.http import JsonResponse
from django.db.models import Q
from .models import EventoCalendario

def autocompletar_evento(request):
    if 'term' in request.GET:
        term = request.GET.get('term')
        eventos = EventoCalendario.objects.filter(
            Q(titulo__icontains=term)
        ).values_list('titulo', flat=True).distinct()
        
        return JsonResponse(list(eventos), safe=False)

def editar_evento(request, evento_id):
    evento = get_object_or_404(EventoCalendario, id=evento_id)
    if request.method == 'POST':
        form = EventoCalendarioForm(request.POST, instance=evento)
        if form.is_valid():
            form.save()
            messages.success(request, "Evento actualizado exitosamente.")
            return redirect('calendario')
    else:
        form = EventoCalendarioForm(instance=evento)
    return render(request, 'calendario/evento_form.html', {'form': form, 'accion': 'Editar Evento'})

def eliminar_evento(request, evento_id):
    evento = get_object_or_404(EventoCalendario, id=evento_id)
    if request.method == "POST":
        evento.delete()
        messages.success(request, "Evento eliminado exitosamente.")
        return redirect('calendario')
    return render(request, 'calendario/confirmar_eliminacion.html', {'evento': evento})


####################################################################################################
from django.urls import reverse_lazy
from django.contrib import messages
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from .models import Unidad
from .forms import UnidadForm

from django.urls import reverse_lazy
from django.contrib import messages
from django.views.generic import ListView, CreateView
from .models import Unidad
from .forms import UnidadForm

from django.urls import reverse_lazy
from django.contrib import messages
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from .models import Unidad
from .forms import UnidadForm
####################################################################################################
# 📌 Listar unidades
# class UnidadListView(ListView):
#     model = Unidad
#     template_name = "unidades/listar_unidades.html"
#     context_object_name = "unidades"
#     ordering = ['curso', 'numero']
from django.core.paginator import Paginator
from django.db.models import Q
from django.views.generic import ListView
from .models import Unidad, Curso

from django.core.paginator import Paginator
from django.db.models import Q
from django.views.generic import ListView
from .models import Unidad, Curso
from django.db.models import Q

from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger

class UnidadListView(ListView):
    model = Unidad
    template_name = "unidades/listar_unidades.html"
    context_object_name = "unidades"
    paginate_by = 10  # 10 unidades por página

    def get_queryset(self):
        queryset = super().get_queryset().select_related("curso__grado__nivel")  

        # Obtener parámetros de búsqueda y filtrado
        query = self.request.GET.get("q", "").strip()
        curso_id = self.request.GET.get("curso", "")
        grado_id = self.request.GET.get("grado", "")
        nivel_id = self.request.GET.get("nivel", "")
        orden = self.request.GET.get("orden", "curso__grado__nivel__nombre")  # Orden por defecto

        # Validar IDs y aplicar filtros
        filtros = {}
        if curso_id.isdigit():
            filtros["curso_id"] = int(curso_id)
        if grado_id.isdigit():
            filtros["curso__grado_id"] = int(grado_id)
        if nivel_id.isdigit():
            filtros["curso__grado__nivel_id"] = int(nivel_id)
        
        queryset = queryset.filter(**filtros)

        if query:
            queryset = queryset.filter(
                Q(nombre__icontains=query) | 
                Q(numero__icontains=query) |
                Q(curso__nombre__icontains=query) |
                Q(curso__grado__nombre__icontains=query) |
                Q(curso__grado__nivel__nombre__icontains=query)
            )

        # Aplicar ordenamiento dinámico
        opciones_orden = {
            "nombre": "nombre",
            "numero": "numero",
            "curso": "curso__nombre",
            "grado": "curso__grado__nombre",
            "nivel": "curso__grado__nivel__nombre"
        }
        queryset = queryset.order_by(opciones_orden.get(orden, "curso__grado__nivel__nombre"))

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        # Paginación mejorada
        paginator = Paginator(self.get_queryset(), self.paginate_by)
        page = self.request.GET.get("page")
        try:
            unidades_paginadas = paginator.page(page)
        except PageNotAnInteger:
            unidades_paginadas = paginator.page(1)
        except EmptyPage:
            unidades_paginadas = paginator.page(paginator.num_pages)

        context["unidades"] = unidades_paginadas
        context["cursos"] = Curso.objects.all()
        context["grados"] = Grado.objects.all()
        context["niveles"] = Nivel.objects.all()
        context["orden_actual"] = self.request.GET.get("orden", "curso__grado__nivel__nombre")

        return context

    



# class UnidadListView(ListView):
#     model = Unidad
#     template_name = "unidades/listar_unidades.html"
#     context_object_name = "unidades"
#     ordering = ['curso__nombre', 'numero']
#     paginate_by = 10  # Paginación: 10 unidades por página

#     def get_queryset(self):
#         queryset = super().get_queryset().select_related("curso")  # Optimizar consulta
#         query = self.request.GET.get("q", "").strip()
#         curso_id = self.request.GET.get("curso", "")

#         # Aplicar filtros
#         if query:
#             queryset = queryset.filter(
#                 Q(nombre__icontains=query) | 
#                 Q(numero__icontains=query)
#             )
#         if curso_id:
#             queryset = queryset.filter(curso_id=curso_id)

#         return queryset

#     def get_context_data(self, **kwargs):
#         context = super().get_context_data(**kwargs)
#         context["cursos"] = Curso.objects.all()  # Para mostrar cursos en el filtro
#         return context

# from django.core.paginator import Paginator
# from django.shortcuts import render
# from .models import Unidad, Curso

# def lista_unidades(request):
#     query = request.GET.get('q', '')
#     curso_id = request.GET.get('curso', '')

#     # Filtrar unidades
#     unidades = Unidad.objects.all()
#     if query:
#         unidades = unidades.filter(nombre__icontains=query)
#     if curso_id:
#         unidades = unidades.filter(curso_id=curso_id)

#     # Paginación
#     paginator = Paginator(unidades, 10)  # 10 unidades por página
#     page_number = request.GET.get('page')
#     unidades_page = paginator.get_page(page_number)

#     # Obtener lista de cursos para los filtros
#     cursos = Curso.objects.all()

#     return render(request, 'unidades/listar_unidades.html', {
#         'unidades': unidades_page,
#         'cursos': cursos,
#         'request': request
#     })
#

# 📌 Crear una nueva unidad
from django.contrib import messages
from django.urls import reverse_lazy
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import CreateView
from .models import Unidad
from .forms import UnidadForm

class UnidadCreateView(LoginRequiredMixin, CreateView):
    model = Unidad
    form_class = UnidadForm
    template_name = "unidades/form_unidad.html"

    def get_success_url(self):
        return reverse_lazy("unidad_list")  # Cambia si usas namespace

    def form_valid(self, form):
        messages.success(self.request, "✅ Unidad creada con éxito.")
        return super().form_valid(form)

    def form_invalid(self, form):
        messages.error(self.request, "❌ Hubo un error al crear la unidad. Verifica los campos.")
        return super().form_invalid(form)


# 📌 Editar unidad
class UnidadUpdateView(UpdateView):
    model = Unidad
    form_class = UnidadForm
    template_name = "unidades/form_unidad.html"

    def get_success_url(self):
        return reverse_lazy("unidad_list")

    def form_valid(self, form):
        messages.success(self.request, "✏️ Unidad actualizada con éxito.")
        return super().form_valid(form)

# 📌 Eliminar unidad
class UnidadDeleteView(DeleteView):
    model = Unidad
    template_name = "unidades/confirmar_eliminar.html"

    def get_success_url(self):
        messages.success(self.request, "🗑️ Unidad eliminada con éxito.")
        return reverse_lazy("unidad_list")


from django.views.generic import DetailView  # ✅ Importa DetailView

# 📌 Ver detalles de una unidad
# 📌 Ver detalles de una unidad
class UnidadDetailView(DetailView):
    model = Unidad
    template_name = "unidades/detalle_unidad.html"
    context_object_name = "unidad"


from django.http import HttpResponse
from django.template.loader import render_to_string
from docx import Document
from docx.shared import Pt
from io import BytesIO
from .models import Unidad

from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from io import BytesIO
from docx import Document
from bs4 import BeautifulSoup  # Para limpiar HTML y convertirlo a texto plano
import html2text
from .models import Unidad  # Asegúrate de importar tu modelo

def exportar_unidad_word(request, pk):
    unidad = get_object_or_404(Unidad, pk=pk)

    # Crear documento Word
    doc = Document()
    doc.add_heading('📌 Detalles de la Unidad', level=1)

    # Nombre de la unidad y curso
    doc.add_paragraph(f'📚 Curso: {unidad.curso.nombre}', style='ListBullet')
    doc.add_paragraph(f'🔢 Número de unidad: {unidad.numero}', style='ListBullet')

    # Convertir HTML a texto enriquecido
    doc.add_paragraph("\n📝 Descripción:", style='Heading2')
    
    # Limpiar el HTML usando BeautifulSoup y html2text
    html_content = unidad.descripcion
    soup = BeautifulSoup(html_content, "html.parser")
    plain_text = html2text.html2text(str(soup))  # Convierte HTML a texto formateado

    doc.add_paragraph(plain_text, style='Normal')

    # Guardar en un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Configurar la respuesta HTTP
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="Unidad_{unidad.nombre}.docx"'
    
    return response
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from bs4 import BeautifulSoup  # Para limpiar HTML y convertirlo a texto plano
import io
from .models import Unidad

from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
from bs4 import BeautifulSoup  # Para limpiar HTML y convertirlo a texto plano
import io
import textwrap
from .models import Unidad

def exportar_unidad_pdf(request, pk):
    unidad = get_object_or_404(Unidad, pk=pk)

    # Crear un buffer de memoria para el PDF
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    # Configurar título
    p.setFont("Helvetica-Bold", 14)
    p.drawString(100, height - 50, "📌 Detalles de la Unidad")

    # Datos de la unidad
    p.setFont("Helvetica", 12)
    p.drawString(100, height - 80, f"📚 Curso: {unidad.curso.nombre}")
    p.drawString(100, height - 100, f"🔢 Número de unidad: {unidad.numero}")

    # Convertir HTML a texto limpio
    p.setFont("Helvetica-Bold", 12)
    p.drawString(100, height - 130, "📝 Descripción:")

    p.setFont("Helvetica", 10)
    html_content = unidad.descripcion
    soup = BeautifulSoup(html_content, "html.parser")
    plain_text = soup.get_text(separator=" ")  # Convertimos HTML a texto limpio sin saltos extraños

    # Dividir la descripción en líneas de máximo 90 caracteres
    wrapped_text = textwrap.wrap(plain_text, width=90)
    
    # Escribir las líneas en el PDF
    y_position = height - 150
    for line in wrapped_text:
        if y_position < 50:  # Si se acerca al final de la página, agrega una nueva
            p.showPage()
            p.setFont("Helvetica", 10)
            y_position = height - 50
        p.drawString(100, y_position, line)
        y_position -= 15  # Espaciado entre líneas

    # Guardar PDF
    p.save()
    buffer.seek(0)

    # Configurar la respuesta HTTP
    response = HttpResponse(buffer, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="Unidad_{unidad.nombre}.pdf"'


    return response
####################################################################################################TEAMAS

from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse_lazy
from django.contrib import messages
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from .models import Tema
from .forms import TemaForm

# ✅ LISTAR TEMAS
class TemaListView(ListView):
    model = Tema
    template_name = "temas/lista.html"
    context_object_name = "temas"
    paginate_by = 10  # Paginación opcional

# ✅ DETALLE DE UN TEMA
class TemaDetailView(DetailView):
    model = Tema
    template_name = "temas/detalle.html"
    context_object_name = "tema"

# ✅ CREAR UN NUEVO TEMA
class TemaCreateView(CreateView):
    model = Tema
    form_class = TemaForm
    template_name = "temas/formulario.html"
    success_url = reverse_lazy("tema_list")

    def form_valid(self, form):
        messages.success(self.request, "✅ Tema creado con éxito.")
        return super().form_valid(form)

# ✅ EDITAR UN TEMA
class TemaUpdateView(UpdateView):
    model = Tema
    form_class = TemaForm
    template_name = "temas/formulario.html"
    success_url = reverse_lazy("tema_list")

    def form_valid(self, form):
        messages.success(self.request, "✅ Tema actualizado con éxito.")
        return super().form_valid(form)

# ✅ ELIMINAR UN TEMA
class TemaDeleteView(DeleteView):
    model = Tema
    template_name = "temas/confirmar_eliminar.html"
    success_url = reverse_lazy("tema_list")

    def delete(self, request, *args, **kwargs):
        messages.success(self.request, "✅ Tema eliminado con éxito.")
        return super().delete(request, *args, **kwargs)

####################################################################################################

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from .models import HorarioDocente
from .forms import HorarioDocenteForm

from django.shortcuts import render
from django.db.models import Q
from .models import HorarioDocente
from .forms import HorarioDocenteForm  # Si necesitas reutilizarlo

from django.shortcuts import render
from .models import HorarioDocente, Docente, CursoAsignado

from django.shortcuts import render
from django.http import JsonResponse
from .models import HorarioDocente, Docente, CursoAsignado

def listar_horarios(request):
    docente_id = request.GET.get('docente', '')
    curso_id = request.GET.get('curso_asignado', '')
    dia = request.GET.get('dia', '')
    hora_inicio = request.GET.get('hora_inicio', '')
    hora_fin = request.GET.get('hora_fin', '')

    horarios = HorarioDocente.objects.select_related('docente', 'curso_asignado').order_by('dia', 'hora_inicio')

    if docente_id:
        horarios = horarios.filter(docente_id=docente_id)
    if curso_id:
        horarios = horarios.filter(curso_asignado_id=curso_id)
    if dia:
        horarios = horarios.filter(dia=dia)
    if hora_inicio:
        horarios = horarios.filter(hora_inicio__gte=hora_inicio)
    if hora_fin:
        horarios = horarios.filter(hora_fin__lte=hora_fin)

    # Si la solicitud es AJAX, devolver solo la tabla
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return render(request, 'horarios/tabla_horarios.html', {'horarios': horarios})

    docentes = Docente.objects.all()
    cursos = CursoAsignado.objects.all()
    dias_semana = HorarioDocente.DIAS_SEMANA

    return render(request, 'horarios/listar_horarios.html', {
        'horarios': horarios,
        'docentes': docentes,
        'cursos': cursos,
        'dias_semana': dias_semana,
    })



from django.shortcuts import render, redirect
from django.contrib import messages
from django.db.models import Q
from .forms import HorarioDocenteForm
from .models import HorarioDocente

def agregar_horario(request):
    if request.method == 'POST':
        form = HorarioDocenteForm(request.POST)
        if form.is_valid():
            docente = form.cleaned_data['docente']
            dia = form.cleaned_data['dia']
            hora_inicio = form.cleaned_data['hora_inicio']
            hora_fin = form.cleaned_data['hora_fin']
            aula = form.cleaned_data['aula']

            # Validación 1: Verificar que la hora de inicio sea menor que la de fin
            if hora_inicio >= hora_fin:
                messages.error(request, "La hora de inicio debe ser menor que la hora de fin.")
                return render(request, 'horarios/agregar_horario.html', {'form': form})

            # Validación 2: Verificar que no haya solapamiento con otro horario del mismo docente
            conflicto_docente = HorarioDocente.objects.filter(
                docente=docente,
                dia=dia
            ).filter(
                Q(hora_inicio__lt=hora_fin, hora_fin__gt=hora_inicio)
            )

            if conflicto_docente.exists():
                messages.error(request, "El docente ya tiene un horario asignado en este día y horario.")
                return render(request, 'horarios/agregar_horario.html', {'form': form})

            # Validación 3 (Opcional): Verificar que el aula no esté ocupada en el mismo horario
            conflicto_aula = HorarioDocente.objects.filter(
                aula=aula,
                dia=dia
            ).filter(
                Q(hora_inicio__lt=hora_fin, hora_fin__gt=hora_inicio)
            )

            if conflicto_aula.exists():
                messages.error(request, "El aula seleccionada ya está ocupada en este horario.")
                return render(request, 'horarios/agregar_horario.html', {'form': form})

            # Si pasa todas las validaciones, se guarda el horario
            form.save()
            messages.success(request, 'Horario agregado correctamente.')
            return redirect('listar_horarios')

    else:
        form = HorarioDocenteForm()
    
    return render(request, 'horarios/agregar_horario.html', {'form': form})

def editar_horario(request, horario_id):
    horario = get_object_or_404(HorarioDocente, id=horario_id)
    if request.method == 'POST':
        form = HorarioDocenteForm(request.POST, instance=horario)
        if form.is_valid():
            form.save()
            messages.success(request, 'Horario actualizado correctamente.')
            return redirect('listar_horarios')
    else:
        form = HorarioDocenteForm(instance=horario)
    return render(request, 'horarios/editar_horario.html', {'form': form})

def eliminar_horario(request, horario_id):
    horario = get_object_or_404(HorarioDocente, id=horario_id)
    if request.method == 'POST':
        horario.delete()
        messages.success(request, 'Horario eliminado correctamente.')
        return redirect('listar_horarios')
    return render(request, 'horarios/eliminar_horario.html', {'horario': horario})


####################################################################################################
from django.shortcuts import render
from django.core.paginator import Paginator
from .models import HorarioDocente, Docente
from django.shortcuts import render
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from .models import HorarioDocente, Docente

from django.shortcuts import render
from .models import HorarioDocente, Docente
from django.core.paginator import Paginator

from django.shortcuts import render
from .models import HorarioDocente, Docente
from django.core.paginator import Paginator

from django.shortcuts import render
from .models import HorarioDocente, Docente
from django.core.paginator import Paginator

from django.core.paginator import Paginator, EmptyPage
from django.shortcuts import render
from .models import HorarioDocente, Docente

# from django.shortcuts import render
# from django.core.paginator import Paginator
# from .models import HorarioDocente, Docente

# def ver_horarios(request):
#     # Obtener todos los docentes con horarios asociados (para evitar consultas innecesarias)
#     docentes = Docente.objects.prefetch_related('horariodocente_set').all()

#     # Obtener las horas disponibles de manera dinámica
#     horarios = HorarioDocente.objects.all()
#     horas_disponibles = sorted(set(horario.hora_inicio.strftime("%H:%M") for horario in horarios))

#     # Definir los días de la semana de manera uniforme
#     DIAS_SEMANA = ['Lunes', 'Martes', 'Miércoles', 'Jueves', 'Viernes']

#     # Crear un diccionario para almacenar los horarios por docente
#     horarios_docentes = {}

#     # Iterar sobre los docentes para obtener sus horarios
#     for docente in docentes:
#         horarios_docente = docente.horariodocente_set.order_by('dia', 'hora_inicio')

#         # Crear un diccionario vacío con la estructura necesaria
#         horario_docente_dict = {
#             dia: {hora: None for hora in horas_disponibles}
#             for dia in DIAS_SEMANA
#         }

#         # Rellenar el diccionario con los horarios del docente
#         for horario in horarios_docente:
#             dia = horario.dia.capitalize()  # Asegurar que el día esté en el formato correcto
#             hora_inicio = horario.hora_inicio.strftime("%H:%M")  # Convertir la hora a string

#             # Asignar el curso al horario correspondiente
#             horario_docente_dict[dia][hora_inicio] = (
#                 f"{horario.curso_asignado.curso.nombre} - {horario.curso_asignado.seccion} "
#                 f"(Prof. {docente.nombre} {docente.apellido})"
#             )
        
#         horarios_docentes[docente] = horario_docente_dict

#     # Manejo de paginación
#     page_number = request.GET.get('page', 1)
#     paginator = Paginator(docentes, 10)  # Mostrar 10 docentes por página
#     page_obj = paginator.get_page(page_number)

#     return render(request, 'horarios/ver_horarios.html', {
#         'horarios_docentes': horarios_docentes,
#         'horas_disponibles': horas_disponibles,
#         'page_obj': page_obj,
#     })


# from django.shortcuts import render
# from django.core.paginator import Paginator
# from .models import HorarioDocente, Docente
# from itertools import groupby

# def ver_horarios(request):
#     # Obtener todos los horarios en una sola consulta optimizada
#     horarios = HorarioDocente.objects.select_related('docente', 'curso_asignado__curso').order_by('dia', 'hora_inicio')

#     # Obtener las horas disponibles dinámicamente
#     horas_disponibles = sorted(set(horario.hora_inicio.strftime("%H:%M") for horario in horarios))

#     # Definir los días de la semana correctamente mapeados
#     DIAS_SEMANA = {
#         'lunes': 'Lunes', 'martes': 'Martes', 'miércoles': 'Miércoles',
#         'jueves': 'Jueves', 'viernes': 'Viernes'
#     }

#     # Agrupar los horarios por docente
#     horarios_docentes = {}
#     for docente, horario_grupo in groupby(horarios, key=lambda h: h.docente):
#         # Crear estructura vacía
#         horario_docente_dict = {dia: {hora: None for hora in horas_disponibles} for dia in DIAS_SEMANA.values()}
        
#         # Rellenar con los horarios del docente
#         for horario in horario_grupo:
#             dia = DIAS_SEMANA.get(horario.dia.lower(), horario.dia)  # Normalizar nombre del día
#             hora_inicio = horario.hora_inicio.strftime("%H:%M")
#             horario_docente_dict[dia][hora_inicio] = (
#                 f"{horario.curso_asignado.curso.nombre} - {horario.curso_asignado.seccion} "
#                 f"(Prof. {docente.nombre} {docente.apellido})"
#             )

#         horarios_docentes[docente] = horario_docente_dict

#     # Manejo de paginación
#     docentes = list(horarios_docentes.keys())  # Extraer docentes de los horarios
#     page_number = request.GET.get('page', 1)
#     try:
#         page_number = int(page_number)
#     except ValueError:
#         page_number = 1

#     paginator = Paginator(docentes, 10)
#     page_obj = paginator.get_page(page_number)

#     return render(request, 'horarios/ver_horarios.html', {
#         'horarios_docentes': horarios_docentes,
#         'horas_disponibles': horas_disponibles,
#         'page_obj': page_obj,
#     })


####################################################################################################Codigo Mejorado

from datetime import datetime, timedelta
from itertools import groupby
from django.core.paginator import Paginator
from django.shortcuts import render
from .models import HorarioDocente

def ver_horarios(request):
    # Obtener todos los horarios en una sola consulta optimizada
    horarios = HorarioDocente.objects.select_related('docente', 'curso_asignado__curso').order_by('dia', 'hora_inicio')

    # Obtener todas las horas únicas en la grilla
    horas_disponibles = sorted(set(horario.hora_inicio.strftime("%H:%M") for horario in horarios))

    # Definir los días de la semana correctamente mapeados
    DIAS_SEMANA = {
        'lunes': 'Lunes', 'martes': 'Martes', 'miércoles': 'Miércoles',
        'jueves': 'Jueves', 'viernes': 'Viernes'
    }

    # Agrupar los horarios por docente
    horarios_docentes = {}
    for docente, horario_grupo in groupby(horarios, key=lambda h: h.docente):
        # Crear estructura vacía
        horario_docente_dict = {dia: {hora: None for hora in horas_disponibles} for dia in DIAS_SEMANA.values()}

        # Rellenar con los horarios del docente
        for horario in horario_grupo:
            dia = DIAS_SEMANA.get(horario.dia.lower(), horario.dia)  # Normalizar nombre del día
            hora_inicio = horario.hora_inicio
            hora_fin = horario.hora_fin

            # Asignación del texto a mostrar
            asignacion = (
                f"{horario.curso_asignado.curso.nombre} - {horario.curso_asignado.seccion} "
                f"(Prof. {docente.nombre} {docente.apellido})"
            )

            # Convertir hora_inicio y hora_fin a datetime para hacer cálculos con timedelta
            hora_actual = datetime.combine(datetime.today(), hora_inicio)
            hora_fin_dt = datetime.combine(datetime.today(), hora_fin)

            # Llenar las celdas desde hora_inicio hasta hora_fin
            while hora_actual < hora_fin_dt:
                hora_actual_str = hora_actual.strftime("%H:%M")
                if hora_actual_str in horario_docente_dict[dia]:  # Solo si la hora está en la tabla
                    horario_docente_dict[dia][hora_actual_str] = asignacion
                
                hora_actual += timedelta(hours=1)  # Avanza de hora en hora

        horarios_docentes[docente] = horario_docente_dict

    # Manejo de paginación
    docentes = list(horarios_docentes.keys())  # Extraer docentes de los horarios
    page_number = request.GET.get('page', 1)
    try:
        page_number = int(page_number)
    except ValueError:
        page_number = 1

    paginator = Paginator(docentes, 10)
    page_obj = paginator.get_page(page_number)

    return render(request, 'horarios/ver_horarios.html', {
        'horarios_docentes': horarios_docentes,
        'horas_disponibles': horas_disponibles,
        'page_obj': page_obj,
    })


####################################################################################################Codigo Mejorado
# from datetime import datetime, timedelta
# from itertools import groupby
# from django.core.paginator import Paginator
# from django.shortcuts import render
# from .models import HorarioDocente, Curso

# # Colores para cursos
# COLORES_CURSOS = ["primary", "success", "warning", "danger", "info", "secondary"]

# def ver_horarios(request):
#     # Obtener todos los horarios
#     horarios = HorarioDocente.objects.select_related('docente', 'curso_asignado__curso').order_by('dia', 'hora_inicio')
    
#     # Obtener horas únicas
#     horas_disponibles = sorted(set(horario.hora_inicio.strftime("%H:%M") for horario in horarios))
    
#     # Definir los días
#     DIAS_SEMANA = {
#         'lunes': 'Lunes', 'martes': 'Martes', 'miércoles': 'Miércoles',
#         'jueves': 'Jueves', 'viernes': 'Viernes'
#     }
    
#     # Asignar colores únicos a cada curso
#     cursos = Curso.objects.all()
#     colores_cursos = {curso.id: COLORES_CURSOS[i % len(COLORES_CURSOS)] for i, curso in enumerate(cursos)}
    
#     # Agrupar horarios por docente
#     horarios_docentes = {}
#     for docente, horario_grupo in groupby(horarios, key=lambda h: h.docente):
#         horario_docente_dict = {dia: {hora: None for hora in horas_disponibles} for dia in DIAS_SEMANA.values()}
        
#         for horario in horario_grupo:
#             dia = DIAS_SEMANA.get(horario.dia.lower(), horario.dia)
#             hora_inicio = horario.hora_inicio.strftime("%H:%M")
#             hora_fin = horario.hora_fin.strftime("%H:%M")
            
#             asignacion = (
#                 f"{horario.curso_asignado.curso.nombre} - {horario.curso_asignado.seccion} "
#                 f"(Prof. {docente.nombre} {docente.apellido})"
#             )
            
#             # Determinar color del curso
#             color = colores_cursos.get(horario.curso_asignado.curso.id, "secondary")
            
#             # Guardar asignación con color
#             horario_docente_dict[dia][hora_inicio] = {"asignacion": asignacion, "color": color}
        
#         horarios_docentes[docente] = horario_docente_dict
    
#     # Paginación
#     docentes = list(horarios_docentes.keys())
#     page_number = request.GET.get('page', 1)
#     paginator = Paginator(docentes, 10)
#     page_obj = paginator.get_page(page_number)
    
#     return render(request, 'horarios/ver_horarios.html', {
#         'horarios_docentes': horarios_docentes,
#         'horas_disponibles': horas_disponibles,
#         'page_obj': page_obj,
#     })


from django.shortcuts import render, redirect
from django.contrib import messages
from .forms import HorarioDocenteForm

def agregar_horarioo(request):
    if request.method == 'POST':
        form = HorarioDocenteForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "✅ Horario agregado correctamente.")
            return redirect('agregar_horario')  # Redirigir a la misma vista después de guardar
        else:
            messages.error(request, "⚠️ Hubo un error en el formulario. Revisa los datos.")

    else:
        form = HorarioDocenteForm()

    return render(request, 'horarios/agregar_horario1.html', {'form': form})

####################################################################################################

from django.shortcuts import render, redirect
from .forms import RegistroForm
from django.contrib.auth.models import Group

from django.contrib.auth.models import User, Group

from django.contrib.auth.models import User, Group

from django.contrib.auth.models import User, Group

def registro_usuario(request):
    if request.method == 'POST':
        form = RegistroForm(request.POST)
        if form.is_valid():
            usuario = form.save(commit=False)
            usuario.set_password(form.cleaned_data['password'])
            usuario.save()  # Se guarda el usuario para obtener un ID

            # Verificar si es el primer usuario
            if User.objects.count() == 1:  # Solo el primero
                usuario.is_superuser = True
                usuario.is_staff = True
                usuario.save()  # Guardar cambios en permisos

                # Asignar al grupo "Administrador"
                grupo_admin, _ = Group.objects.get_or_create(name='Administrador')
                usuario.groups.add(grupo_admin)
            else:
                # Asignar al grupo "Profesor" a los demás usuarios
                grupo_profesor, _ = Group.objects.get_or_create(name='Profesor')
                usuario.groups.add(grupo_profesor)

            return redirect('login')
    else:
        form = RegistroForm()

    return render(request, 'registration/registro.html', {'form': form})





from django.shortcuts import render
from django.contrib.auth.decorators import login_required

@login_required
def home(request):
    return render(request, 'home.html')
from django.shortcuts import render


from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.models import Group, User
from django.contrib import messages
from .decorators import administrador_requerido

@administrador_requerido  # Restringe el acceso solo a los Administradores
def asignar_usuarios_a_grupo(request, group_id):
    grupo = get_object_or_404(Group, id=group_id)
    usuarios = User.objects.all()

    if request.method == 'POST':
        usuarios_seleccionados = request.POST.getlist('usuarios')
        grupo.user_set.set(usuarios_seleccionados)
        messages.success(request, f'Usuarios asignados al grupo {grupo.name} correctamente.')
        return redirect('group_list')

    return render(request, 'groups/asignar_usuarios.html', {'grupo': grupo, 'usuarios': usuarios})



def salida(request):
    return render(request, 'registration/salida.html')

####################################################################################################Detalle del Alumno






####################################################################################################USUARIOS
from django.contrib.auth.models import User
# Lista de usuarios
class UserListView(ListView):
    model = User
    template_name = "usuarios/user_list.html"
    context_object_name = "usuarios"

# Editar usuario
class UserUpdateView(UpdateView):
    model = User
    form_class = UserForm
    template_name = "usuarios/user_form.html"
    success_url = reverse_lazy("user_list")

# Eliminar usuario
class UserDeleteView(DeleteView):
    model = User
    template_name = "usuarios/user_confirm_delete.html"
    success_url = reverse_lazy("user_list")